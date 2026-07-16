"""Text extraction from uploaded student documents (PDF or DOCX)."""

from __future__ import annotations

import io

MAX_CHARS = 180_000  # keep prompts within reach of all supported providers


class ExtractionError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    else:
        raise ExtractionError("Unsupported file type. Please upload a .pdf or .docx file.")

    text = text.strip()
    if len(text) < 200:
        raise ExtractionError(
            "Could not extract enough text from the document. "
            "If this is a scanned PDF, it needs OCR before it can be checked."
        )
    if len(text) > MAX_CHARS:
        # Keep the start (body) and the end (reference list) — the middle is
        # least likely to contain the bibliography.
        head = int(MAX_CHARS * 0.6)
        tail = MAX_CHARS - head
        text = text[:head] + "\n\n[... document truncated ...]\n\n" + text[-tail:]
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # pypdf raises many exception types
        raise ExtractionError(f"Failed to read PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Failed to read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
